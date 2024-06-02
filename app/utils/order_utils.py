import os
import shutil
from docx import Document
from datetime import datetime
import base64
import aiofiles

async def create_contract(file_name, replacements):
    # Define file paths
    original_file = 'files/contract.docx'
    contract_dir = 'files/contract'
    timestamp = datetime.now().timestamp()
    file_name += f"_{timestamp}"
    copied_file = os.path.join(contract_dir, f'{file_name}.docx')

    # Create the contract directory if it doesn't exist
    os.makedirs(contract_dir, exist_ok=True)

    # Copy the original file to the contract directory
    shutil.copy2(original_file, copied_file)
    # Load the copied Word document
    doc = Document(copied_file)
    # doc.save(copied_file)

    # Iterate through paragraphs and replace text
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)

    # Iterate through tables and replace text in cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old_text, new_text in replacements.items():
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_text)

    # Save the modified document
    doc.save(copied_file)
    return copied_file

async def base64_to_pdf(base64_string, filename):
    # Clean up the Base64 string (if needed)
    base64_string = base64_string.replace("\n", "")

    # Decode the Base64 string
    pdf_data = base64.b64decode(base64_string)

    # Create the contract directory if it doesn't exist
    contract_dir = 'files/contract'
    os.makedirs(contract_dir, exist_ok=True)
    # Specify the output file path
    timestamp = int(datetime.now().timestamp())
    output_file_path = f"{contract_dir}/{filename}_{timestamp}.pdf"
    # Write the binary data to a PDF file
    async with aiofiles.open(output_file_path, 'wb') as pdf_file:
        await pdf_file.write(pdf_data)
    return output_file_path